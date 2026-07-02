"""Gerador de PGR.docx — Sprint 4.

Recebe o mesmo JSON que o Motor de Decisão produz (montar_inventario_empresa)
e gera o documento. Na versão gratuita, aplica marca d'água como IMAGEM no
cabeçalho de cada página (não é texto — não dá pra selecionar/apagar editando
o corpo do documento).
"""
import io
import os
from datetime import date
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont

WATERMARK_TEXT = "DEMONSTRAÇÃO — VERSÃO NÃO LICENCIADA"


def _gerar_imagem_marca_dagua() -> io.BytesIO:
    """Gera um PNG com texto diagonal semi-transparente para usar como marca d'água."""
    img = Image.new("RGBA", (1600, 1600), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 60)
    except Exception:
        font = ImageFont.load_default()

    txt_img = Image.new("RGBA", (1400, 150), (255, 255, 255, 0))
    txt_draw = ImageDraw.Draw(txt_img)
    txt_draw.text((0, 0), WATERMARK_TEXT, font=font, fill=(200, 30, 30, 90))
    txt_img = txt_img.rotate(45, expand=True)
    img.paste(txt_img, (100, 700), txt_img)

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


def _inserir_marca_dagua_no_cabecalho(document: Document) -> None:
    """Insere a imagem da marca d'água no header de todas as seções — imagem,
    não texto, para não ser removível editando o corpo do documento."""
    watermark_buf = _gerar_imagem_marca_dagua()
    for section in document.sections:
        header = section.header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(watermark_buf, width=Inches(6.5))
        watermark_buf.seek(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Posiciona a imagem atrás do texto do corpo (behind text) via XML direto
        inline = paragraph.runs[-1]._element.findall(qn("w:drawing"))
        for drawing in inline:
            anchor = OxmlElement("wp:anchor")
            # Mantém simples: a imagem no header já fica atrás do conteúdo do
            # corpo por padrão em headers do Word (não sobrepõe a leitura).


def gerar_pgr_docx(inventario: dict, empresa: dict, licenca_premium: bool = False) -> io.BytesIO:
    """Gera o documento PGR completo a partir do resultado do Motor de Decisão."""
    document = Document()

    if not licenca_premium:
        _inserir_marca_dagua_no_cabecalho(document)

    titulo = document.add_heading("PGR — Programa de Gerenciamento de Riscos", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph(f"Empresa: {empresa.get('razao_social', '')}")
    document.add_paragraph(f"CNPJ: {empresa.get('cnpj', '')}")
    document.add_paragraph(f"Data de emissão: {date.today().strftime('%d/%m/%Y')}")
    if not licenca_premium:
        p = document.add_paragraph()
        run = p.add_run("VERSÃO DE DEMONSTRAÇÃO — documento não válido para fins legais até licenciamento.")
        run.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0x1E, 0x1E)

    document.add_heading("1. Inventário de Riscos", level=1)

    for item in inventario.get("inventario", []):
        atividade = item["atividade"]
        document.add_heading(f"Atividade: {atividade['nome']}", level=2)

        if item.get("aviso"):
            document.add_paragraph(item["aviso"])
            continue

        for perigo in item.get("perigos", []):
            p = document.add_paragraph()
            p.add_run(f"Perigo: {perigo['nome']} ({perigo['categoria']})").bold = True
            document.add_paragraph(f"Descrição: {perigo.get('descricao', '')}")
            document.add_paragraph(f"Consequências: {perigo.get('consequencias', '')}")
            document.add_paragraph(f"Base legal: {perigo.get('base_legal', '')}")

            if perigo.get("controles"):
                document.add_paragraph("Medidas de Controle:")
                for c in perigo["controles"]:
                    document.add_paragraph(f"  • [{c['tipo']}] {c['nome']}", style="List Bullet")

            if perigo.get("epis"):
                document.add_paragraph("EPIs indicados: " + ", ".join(e["nome"] for e in perigo["epis"]))

            if perigo.get("epcs"):
                document.add_paragraph("EPCs indicados: " + ", ".join(e["nome"] for e in perigo["epcs"]))

            if perigo.get("treinamentos"):
                document.add_paragraph("Treinamentos exigidos: " + ", ".join(t["nome"] for t in perigo["treinamentos"]))

    document.add_heading("2. Documentos Complementares Necessários", level=1)
    for doc_tipo in inventario.get("documentos_necessarios", []):
        document.add_paragraph(f"• {doc_tipo}", style="List Bullet")

    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf
